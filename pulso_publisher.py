#!/usr/bin/env python3
"""
Genera PULSO a la IA — Edición 23
24 de mayo de 2026 | EMPRENDEDORES.LTD
Contenido curado manualmente con noticias reales de la semana.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from datetime import datetime

OUTPUT_DIR = Path(__file__).resolve().parent / "output"
OUTPUT_DIR.mkdir(exist_ok=True)

EDITION = 23
DATE_STR = "24 de mayo de 2026"
AZUL = RGBColor(0x00, 0x33, 0x66)      # Azul EMPRENDEDORES.LTD
AZUL_CLARO = RGBColor(0x00, 0x70, 0xC0)
GRIS = RGBColor(0x60, 0x60, 0x60)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)

# ── Datos editoriales ──────────────────────────────────────────────────────────

RESUMEN_EJECUTIVO = [
    "La semana del 18 al 24 de mayo de 2026 marca un punto de inflexión en la adopción de IA en el sector financiero latinoamericano. El evento AIFI26 en Buenos Aires reunió a más de 1.500 ejecutivos bancarios y fintech bajo una sola tesis: la economía de agentes de IA ha llegado para quedarse. El 30% de los casos de uso globales de IA ya son agénticos, y el sector financiero lidera esa transición.",
    "Al mismo tiempo, Google reestructuró sus planes Gemini con el lanzamiento de Gemini 3.5 Flash como modelo de referencia para automatización empresarial, mientras Alteryx presentó Agent Studio con soporte MCP, democratizando la creación de agentes para analistas de negocio sin necesidad de equipos de IT. En Ecuador, la Estrategia Nacional de IA 2025-2029 avanza junto con un proyecto de Ley Orgánica en la Asamblea Nacional.",
    "Para el ejecutivo latinoamericano de banca, agroindustria y comercio, la pregunta ya no es si adoptar IA agéntica — sino cuándo y con qué herramienta. Esta edición entrega el mapa de la semana con las señales más accionables."
]

NOTICIAS = [
    {
        "titulo": "AIFI26: La banca latinoamericana entra en la era de los agentes de IA",
        "que_paso": "El 21 de mayo, Buenos Aires fue sede del evento AI in Finance 2026 (AIFI26), con más de 1.500 asistentes y CEOs de Banco Galicia, Macro, Santander y Supervielle. El foco fue la transición de IA operativa a IA agéntica — sistemas que no solo automatizan tareas sino que ejecutan decisiones.",
        "por_que_importa": "El 30% de los casos de uso de IA globales ya son agénticos (Q1 2026). Banco de Bogotá, Davivienda e Interbank están pilotando agentes en WhatsApp para onboarding completo, KYC y modelos de riesgo. El sector 'agentic e-commerce' podría generar hasta USD 1,5 billones en volumen económico en los próximos años.",
        "te_afecta": "Si diriges operaciones en banca, seguros o agroindustria, los agentes autónomos están dejando de ser I+D para convertirse en ventaja competitiva real. Los bancos que piloten hoy definirán los estándares de la industria regional en 2027.",
        "recomendacion": "Identifica un proceso transaccional repetitivo (apertura de cuenta, validación documental, atención postventa) y arma un piloto agéntico en 60 días. AIFI26 confirma que el tiempo de experimentación cerró.",
        "url": "https://www.elobservador.com.uy/argentina/economia-y-negocios/aifi-2026-bancos-y-fintechs-aceleran-la-economia-agentes-ia-n6044927"
    },
    {
        "titulo": "Ecuador aprueba Estrategia Nacional de IA 2025-2029 y avanza hacia Ley Orgánica",
        "que_paso": "El Ministerio de Telecomunicaciones de Ecuador emitió en enero de 2026 la Estrategia EFIA-EC, hoja de ruta nacional para adopción ética y responsable de IA hasta 2029. Paralelamente, un Proyecto de Ley Orgánica de Regulación y Promoción de la Inteligencia Artificial está siendo procesado en la Asamblea Nacional.",
        "por_que_importa": "Ecuador pasa de 'adoptante tardío' (según el Índice ILIA 2025) a tener marco regulatorio formal. La ley traería obligaciones concretas para empresas que usen IA en servicios financieros, salud y servicios públicos. En paralelo, más de 150 proyectos de ley de IA están activos en toda Latinoamérica.",
        "te_afecta": "Empresas con operaciones en Ecuador que ya usen IA para decisiones crediticias, scoring o atención al cliente deben mapear sus sistemas frente a los tres ejes de la EFIA-EC: gobernanza, capacidades tecnológicas y desarrollo ético.",
        "recomendacion": "Asigna a Legal/Compliance la tarea de revisar el texto del Proyecto de Ley Orgánica en la Asamblea Nacional. Documenta ya qué sistemas usan IA en tu operación ecuatoriana — es más fácil cumplir ahora que remediar después.",
        "url": "https://www.lexis.com.ec/noticias/registro-oficial-del-dia-ecuador-aprueba-estrategia-nacional-para-el-desarrollo-y-uso-etico-de-la-inteligencia-artificial"
    },
    {
        "titulo": "Alteryx lanza Agent Studio y servidor MCP: los analistas construyen los agentes",
        "que_paso": "El 20 de mayo, en la conferencia Inspire 2026, Alteryx presentó Agent Studio y un servidor MCP (Model Context Protocol). Permite a analistas de negocio convertir flujos de datos y lógica existente en agentes autónomos sin depender de IT. Los agentes se integran con Slack, Microsoft Teams, Claude y OpenAI.",
        "por_que_importa": "Marca el inicio del paradigma 'analista como constructor de agentes'. Las organizaciones con datos bien gobernados pueden operacionalizar IA sin contratar ingenieros especializados. Los flujos se versionan automáticamente y tienen gobernanza y aprobación incorporadas.",
        "te_afecta": "Tus analistas financieros o de operaciones ya conocen los datos del negocio mejor que cualquier equipo externo. Con herramientas como Alteryx Agent Studio, pueden transformar reportes manuales y procesos de conciliación en agentes que corren solos.",
        "recomendacion": "Haz un inventario de los 5 informes o procesos de datos que más tiempo consumen en tu área. Evalúa si Alteryx One puede convertir alguno en agente autónomo esta semana — la prueba gratuita está disponible.",
        "url": "https://www.prnewswire.com/news-releases/alteryx-puts-business-logic-at-the-center-of-agentic-ai-enabling-enterprises-to-operationalize-ai-at-scale-302776782.html"
    }
]

MODELOS = [
    {
        "titulo": "Gemini 3.5 Flash — Google I/O 2026",
        "descripcion": "Google lanzó Gemini 3.5 Flash como modelo predeterminado para tareas complejas de programación y automatización empresarial. Opera a 4x la velocidad de modelos comparables, con un contexto de 1 millón de tokens y precios de $1.50/$9 por millón de tokens (entrada/salida). Google también presentó el plan AI Ultra a $100/mes y Gemini Omni para creación de contenido multimedia.",
        "recomendacion": "Ideal para pilotos corporativos de análisis de documentos extensos (contratos, estados financieros, auditorías) donde la velocidad y el costo por token son variables críticas. El contexto de 1M tokens permite procesar libros enteros de cuentas en una sola llamada."
    },
    {
        "titulo": "GPT-5.5 Instant vs Claude Opus 4.7 Thinking — Ranking mayo 2026",
        "descripcion": "Según el benchmark Chatbot Arena de mayo 2026, GPT-5.5 lidera con 1.506 puntos Elo, seguido de cerca por Claude Opus 4.7 Thinking (1.505, líder en codificación pura) y Gemini 3.1 Pro (1.505). Apple anunció que iOS 27, iPadOS 27 y macOS 27 permitirán a los usuarios elegir entre los tres modelos como motor de Apple Intelligence.",
        "recomendacion": "Para equipos técnicos: Claude Opus 4.7 Thinking es la mejor opción para generación de código y análisis complejo. Para tareas conversacionales y de productividad general en el ecosistema Apple, esperar iOS 27 para evaluar en contexto empresarial."
    }
]

TENDENCIAS = [
    {
        "titulo": "77% de fintechs mexicanas ya usan IA — el umbral de entrada subió",
        "analisis": "El Finnovista Fintech Radar México 2026 revela que 77% de las fintechs en México usan IA y 27% operan bajo modelos AI-first. El 45.2% de las fintechs de pagos y remesas ya aplican IA para monitoreo de fraude en tiempo real. La implicación para el ejecutivo regional es directa: la IA operacional dejó de ser diferenciador para convertirse en requisito de entrada al mercado. Las empresas financieras que no cuenten con al menos detección de fraude por IA en 2026 enfrentan desventaja competitiva estructural."
    },
    {
        "titulo": "Más de 150 proyectos de ley de IA activos en Latinoamérica: la ventana regulatoria se cierra",
        "analisis": "Según Niubox Legal y Polifonía, más de 150 proyectos de ley relacionados con IA están activos en la región. México, Argentina y Perú lideran la iniciativa legislativa, con el marco europeo de escalas de riesgo como modelo predominante. Para 2026-2030 se esperan leyes aprobadas en al menos 6 países. El período actual es la última ventana para que las empresas adopten IA con flexibilidad — en 18-24 meses, las obligaciones formales de transparencia, auditoría y explicabilidad estarán vigentes en varios mercados latinoamericanos."
    }
]

# ── Hashtags ──────────────────────────────────────────────────────────────────
# Fijos (van en config.yaml en el pipeline automático)
HASHTAGS_FIJOS = [
    "#PulsoALaIA", "#InteligenciaArtificial", "#IA",
    "#EMPRENDEDORES", "#Newsletter", "#Tecnologia", "#LiderazgoDigital",
]
# Dinámicos — generados según el contenido de esta edición
HASHTAGS_DINAMICOS = [
    "#AgentesIA", "#AIFI26", "#BancaDigital", "#FinTech",
    "#EcuadorIA", "#RegulacionIA", "#Gemini35Flash",
    "#AlteryxMCP", "#AutomatizacionEmpresarial", "#LatinoAmerica",
]

VEREDICTO = [
    "Evalúa en los próximos 30 días qué proceso repetitivo puede ser tu primer piloto agéntico — onboarding, conciliación de estados de cuenta o atención postventa son los puntos de entrada más probados en banca latinoamericana.",
    "Prueba Gemini 3.5 Flash para análisis de contratos o documentos financieros extensos: su precio por token y contexto de 1M lo hacen el más accesible para piloto corporativo sin presupuesto dedicado.",
    "Asigna a un analista de negocio (no IT) para explorar Alteryx Agent Studio esta semana — la tendencia global es que quienes conocen los datos del negocio sean quienes construyan los agentes.",
    "Si operas en Ecuador, revisa tu posición frente a la Estrategia EFIA-EC 2025-2029 y el Proyecto de Ley Orgánica en la Asamblea; el cumplimiento preventivo cuesta 10x menos que la remediación posterior.",
    "Suscríbete al Índice Latinoamericano de IA (ILIA) para monitorear cómo escalan las regulaciones país por país — es la herramienta de referencia para anticipar obligaciones legales antes de que sean exigibles."
]


# ── Utilidades de formato ──────────────────────────────────────────────────────

def truncate_text(text, max_len=55):
    if not text:
        return ""
    if len(text) <= max_len:
        return text
    truncated = text[:max_len]
    last_space = truncated.rfind(' ')
    if last_space > 0:
        truncated = truncated[:last_space]
    return truncated.strip() + "..."


def format_list_index(titles, max_len=50):
    if not titles:
        return ""
    if len(titles) == 1:
        return truncate_text(titles[0], max_len)
    elif len(titles) == 2:
        joined = f"{titles[0]} y {titles[1]}"
        if len(joined) <= max_len:
            return joined
    first = truncate_text(titles[0], max_len - 12)
    return f"{first} (y {len(titles)-1} más)"


class LayoutTracker:
    def __init__(self, page_height_limit=620):  # Alto útil real de página carta (puntos)
        self.current_height = 0
        self.limit = 310  # Límite del 50% de la página (la mitad)

    def add_height(self, points):
        if self.current_height + points <= 620:
            self.current_height += points
        else:
            # Word empuja el bloque entero a la siguiente página.
            # Por lo tanto, la altura en la nueva página empieza exactamente con el tamaño de este bloque.
            self.current_height = points

    def force_page_break(self):
        self.current_height = 0

    def check_break_before_section(self, doc, section_height=130):
        # Desactivado a petición del usuario para dejar que Word maneje el flujo libremente.
        return False


def add_colored_paragraph(doc, text, color=AZUL, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.color.rgb = color
    run.font.size = Pt(size)
    return p


def add_section_header(doc, text, space_before=14, space_after=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = AZUL
    # Línea decorativa bajo el header
    border = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '003366')
    border.append(bottom)
    p._p.get_or_add_pPr().append(border)
    return p


def add_label_value(doc, label, value, label_color=AZUL_CLARO, value_color=RGBColor(0x1a,0x1a,0x1a)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(label + " ")
    r1.bold = True
    r1.font.color.rgb = label_color
    r1.font.size = Pt(10)
    r2 = p.add_run(value)
    r2.font.color.rgb = value_color
    r2.font.size = Pt(10)
    return p


def set_margins(doc, top=2.5, bottom=2.2, left=2.5, right=2.5):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)


def add_run_to_para(para, text, bold=False, size=9, color=GRIS, italic=False):
    r = para.add_run(text)
    r.bold   = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return r


def add_page_number(para):
    """Inserta campo PAGE en el párrafo dado (field chars dentro de w:r)."""
    r_begin = OxmlElement('w:r')
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    r_begin.append(fld_begin)
    para._p.append(r_begin)

    r_instr = OxmlElement('w:r')
    instr = OxmlElement('w:instrText')
    instr.text = ' PAGE '
    r_instr.append(instr)
    para._p.append(r_instr)

    r_end = OxmlElement('w:r')
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r_end.append(fld_end)
    para._p.append(r_end)


def build_footer_para(para):
    """Construye el footer estándar en el párrafo dado."""
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = para._p.get_or_add_pPr()

    # Borde superior azul
    pBdr = OxmlElement('w:pBdr')
    top_b = OxmlElement('w:top')
    top_b.set(qn('w:val'), 'single')
    top_b.set(qn('w:sz'), '4')
    top_b.set(qn('w:space'), '4')
    top_b.set(qn('w:color'), '003366')
    pBdr.append(top_b)
    pPr.append(pBdr)

    # Tab stops: centro y derecha
    tabs = OxmlElement('w:tabs')
    tc = OxmlElement('w:tab')
    tc.set(qn('w:val'), 'center')
    tc.set(qn('w:pos'), '4536')
    tabs.append(tc)
    tr = OxmlElement('w:tab')
    tr.set(qn('w:val'), 'right')
    tr.set(qn('w:pos'), '9072')
    tabs.append(tr)
    pPr.append(tabs)

    # Izquierda
    r1 = para.add_run("emprendedores.ec  |  ricardowec@gmail.com")
    r1.font.size = Pt(8)
    r1.font.color.rgb = GRIS

    # Centro: tab + "Pág. N"
    tab1 = OxmlElement('w:r')
    tab1.append(OxmlElement('w:tab'))
    para._p.append(tab1)
    r_pn = para.add_run("Pág. ")
    r_pn.font.size = Pt(8)
    r_pn.font.color.rgb = GRIS
    add_page_number(para)

    # Derecha: tab + suscripción
    tab2 = OxmlElement('w:r')
    tab2.append(OxmlElement('w:tab'))
    para._p.append(tab2)
    r_fr = para.add_run("emprendedores.ec/suscripcion")
    r_fr.font.size = Pt(8)
    r_fr.font.color.rgb = AZUL_CLARO


def _remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBdr = OxmlElement('w:tcBdr')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tcBdr.append(b)
    tcPr.append(tcBdr)


def _add_para_bottom_border(para, color='003366'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '4')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def _set_cell_bg(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def _set_cell_margins(cell, top=280, left=280, bottom=280, right=280):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)


def _make_no_border_table(doc, rows=1, cols=1):
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl_xml = tbl._tbl

    tblPr = tbl_xml.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl_xml.insert(0, tblPr)

    # Leer el ancho real de texto que python-docx ya calculó en tblGrid
    tblGrid = tbl_xml.find(qn('w:tblGrid'))
    if tblGrid is not None:
        total_w = sum(int(gc.get(qn('w:w'), 0))
                      for gc in tblGrid.findall(qn('w:gridCol')))
    else:
        # Fallback: calcular desde la sección
        s = doc.sections[0]
        total_w = int((s.page_width - s.left_margin - s.right_margin) / 635)

    # Reemplazar tblW auto/0 con el ancho real del área de texto
    existing = tblPr.find(qn('w:tblW'))
    if existing is not None:
        tblPr.remove(existing)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(total_w))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    # Layout fijo: impide que Word redimensione con autofit
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    # Sin bordes visibles
    tblBdr = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        tblBdr.append(b)
    tblPr.append(tblBdr)
    return tbl


def add_title_block(doc, edition, date_str):
    """Portada: cuadro con fondo azul oscuro, texto en blanco."""
    tbl = _make_no_border_table(doc)
    cell = tbl.cell(0, 0)
    _remove_cell_borders(cell)
    _set_cell_bg(cell, '003366')
    _set_cell_margins(cell, top=300, left=440, bottom=300, right=440)

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(f"PULSO a la IA  ·  Edición {edition}")
    r1.bold = True
    r1.font.size = Pt(26)
    r1.font.color.rgb = BLANCO

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(3)
    r2 = p2.add_run("Inteligencia Artificial para ejecutivos de banca,\ncomercio y agroindustria en Latinoamérica")
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0xCC, 0xE5, 0xFF)

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(2)
    p3.paragraph_format.space_after = Pt(4)
    r3 = p3.add_run(f"EMPRENDEDORES.LTD  ·  {date_str}")
    r3.bold = True
    r3.font.size = Pt(11)
    r3.font.color.rgb = RGBColor(0x66, 0xB8, 0xFF)


def add_summary_block(doc, paragraphs_text):
    """Resumen ejecutivo: cuadro con fondo celeste claro, título incluido."""
    tbl = _make_no_border_table(doc)
    cell = tbl.cell(0, 0)
    _remove_cell_borders(cell)
    _set_cell_bg(cell, 'E3F2FD')
    _set_cell_margins(cell, top=180, left=360, bottom=180, right=360)

    # Título dentro del recuadro
    p_title = cell.paragraphs[0]
    p_title.paragraph_format.space_before = Pt(1)
    p_title.paragraph_format.space_after = Pt(3)
    r_t = p_title.add_run("📊 RESUMEN EJECUTIVO")
    r_t.bold = True
    r_t.font.size = Pt(11)
    r_t.font.color.rgb = AZUL
    _add_para_bottom_border(p_title)

    for texto in paragraphs_text:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = Pt(18)
        r = p.add_run(texto)
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)


def _build_header_table(container):
    """Tabla invisible 2 columnas: izquierda PULSO info, derecha EMPRENDEDORES.LTD."""
    # Limpiar contenido existente
    for p in container.paragraphs:
        p.clear()

    # Crear tabla 1 fila x 2 columnas
    tbl = container.add_table(rows=1, cols=2, width=Cm(16))

    # Quitar bordes de la tabla completa
    tbl_xml = tbl._tbl
    tblPr = tbl_xml.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl_xml.insert(0, tblPr)
    tblBdr = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        tblBdr.append(b)
    tblPr.append(tblBdr)

    # Anchos: 11cm izquierda, 5cm derecha
    from docx.oxml import OxmlElement as OE
    for i, width_cm in enumerate([11, 5]):
        cell = tbl.cell(0, i)
        _remove_cell_borders(cell)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(int(width_cm * 567)))  # cm a twips: 1cm = 567 twips
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

    # Celda izquierda — forzar LEFT via XML (el estilo Header centra por defecto)
    p_l = tbl.cell(0, 0).paragraphs[0]
    pPr_l = p_l._p.get_or_add_pPr()
    jc_l = OxmlElement('w:jc')
    jc_l.set(qn('w:val'), 'left')
    pPr_l.append(jc_l)
    r1 = p_l.add_run("PULSO a la IA")
    r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = AZUL
    r2 = p_l.add_run(f"  ·  Edición {EDITION}  ·  {DATE_STR}")
    r2.font.size = Pt(9); r2.font.color.rgb = GRIS

    # Celda derecha — forzar RIGHT via XML
    p_r = tbl.cell(0, 1).paragraphs[0]
    pPr_r = p_r._p.get_or_add_pPr()
    jc_r = OxmlElement('w:jc')
    jc_r.set(qn('w:val'), 'right')
    pPr_r.append(jc_r)
    r3 = p_r.add_run("EMPRENDEDORES.LTD")
    r3.bold = True; r3.font.size = Pt(10); r3.font.color.rgb = AZUL_CLARO

    # Borde inferior en el párrafo que Word añade tras la tabla
    last_p = container.paragraphs[-1]
    _add_para_bottom_border(last_p)


def setup_header_footer(doc):
    section = doc.sections[0]

    # Activar primera página diferente — vía XML directo (más fiable que la propiedad)
    sectPr = section._sectPr
    if sectPr.find(qn('w:titlePg')) is None:
        titlePg = OxmlElement('w:titlePg')
        sectPr.append(titlePg)

    # ── HEADER PÁGINA 1: vacío (sin header en portada) ───────────────────────
    first_hdr = section.first_page_header
    first_hdr.is_linked_to_previous = False
    for p in first_hdr.paragraphs:
        p.clear()

    # ── FOOTER PÁGINA 1: igual al footer estándar ────────────────────────────
    first_ftr = section.first_page_footer
    first_ftr.is_linked_to_previous = False
    for p in first_ftr.paragraphs:
        p.clear()
    build_footer_para(first_ftr.paragraphs[0])

    # ── HEADER PÁGINAS 2+: tabla izq/der con borde inferior ──────────────────
    _build_header_table(section.header)

    # ── FOOTER PÁGINAS 2+: igual al de página 1 ──────────────────────────────
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()
    build_footer_para(footer.paragraphs[0])


# ── Generación del documento ───────────────────────────────────────────────────

def generate(edition, date_str, resumen_ejecutivo, noticias, modelos, tendencias,
             veredicto, hashtags_fijos, hashtags_dinamicos, output_path):
    doc = Document()
    set_margins(doc)
    setup_header_footer(doc)

    # Eliminar el primer párrafo vacío por defecto de Word
    if len(doc.paragraphs) > 0 and doc.paragraphs[0].text == "":
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    # Fuente cuerpo 12pt — Normal style Y docDefaults
    doc.styles['Normal'].font.size = Pt(12)
    styles_el = doc.styles.element
    doc_defaults = styles_el.find(qn('w:docDefaults'))
    if doc_defaults is not None:
        rPr = doc_defaults.find('.//' + qn('w:rPr'))
        if rPr is not None:
            for tag in [qn('w:sz'), qn('w:szCs')]:
                el = rPr.find(tag)
                if el is not None:
                    el.set(qn('w:val'), '24')  # 24 half-points = 12pt

    # ── PORTADA ───────────────────────────────────────────────────────────────
    add_title_block(doc, edition, date_str)

    # Índice de contenidos
    add_section_header(doc, "En esta edición", space_before=8, space_after=2)
    
    num_acciones = len(veredicto)
    
    resumen_desc = f"Tesis: {truncate_text(resumen_ejecutivo[0], 50)}" if resumen_ejecutivo else "Tesis de la semana"
    noticias_desc = format_list_index([n['titulo'] for n in noticias], 50)
    modelos_desc = format_list_index([m['titulo'] for m in modelos], 50)
    tendencias_desc = format_list_index([t['titulo'] for t in tendencias], 50)
    veredicto_desc = f"{num_acciones} acciones clave para implementar esta semana"
    
    indice = [
        ("📊", "Resumen Ejecutivo", resumen_desc),
        ("📰", "Noticias Destacadas", noticias_desc),
        ("🤖", "Modelos Destacados", modelos_desc),
        ("📈", "Tendencias del Mercado", tendencias_desc),
        ("✅", "Veredicto Accionable", veredicto_desc),
    ]
    for emoji, titulo, desc in indice:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r1 = p.add_run(f"{emoji}  {titulo}: ")
        r1.bold = True
        r1.font.size = Pt(12)
        r1.font.color.rgb = AZUL
        r2 = p.add_run(desc)
        r2.font.size = Pt(12)
        r2.font.color.rgb = GRIS

    # ── RESUMEN EJECUTIVO ─────────────────────────────────────────────────────
    add_summary_block(doc, resumen_ejecutivo)

    # Inicializar tracker de diseño para las siguientes páginas
    tracker = LayoutTracker()

    # ── NOTICIAS DESTACADAS (Páginas 2–4) ────────────────────────────────────
    doc.add_page_break()
    tracker.force_page_break()
    add_section_header(doc, "📰 Noticias Destacadas")
    tracker.add_height(35)
    
    for i, n in enumerate(noticias, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(f"Noticia {i}  —  {n['titulo']}")
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = AZUL
        
        title_lines = 2 if len(n['titulo']) > 55 else 1
        tracker.add_height(17 + title_lines * 15)

        for label, key in [
            ("Qué pasó:", "que_paso"),
            ("Por qué importa:", "por_que_importa"),
            ("Te afecta:", "te_afecta"),
            ("Recomendación:", "recomendacion"),
        ]:
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_before = Pt(4)
            p2.paragraph_format.space_after = Pt(4)
            p2.paragraph_format.line_spacing = Pt(17)
            r1 = p2.add_run(label + "  ")
            r1.bold = True
            r1.font.size = Pt(12)
            r1.font.color.rgb = AZUL_CLARO
            r2 = p2.add_run(n[key])
            r2.font.size = Pt(12)
            r2.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
            
            desc_lines = max(1, len(n[key]) // 65)
            tracker.add_height(8 + desc_lines * 17)

        p3 = doc.add_paragraph()
        p3.paragraph_format.space_before = Pt(3)
        p3.paragraph_format.space_after = Pt(8)
        r3 = p3.add_run(f"Fuente: {n['url']}")
        r3.font.size = Pt(8)
        r3.font.color.rgb = GRIS
        r3.italic = True
        tracker.add_height(11 + 8)

    # ── MODELOS DESTACADOS (Página 4) ────────────────────────────────────────
    tracker.check_break_before_section(doc, section_height=150)
    add_section_header(doc, "🤖 Modelos Destacados")
    tracker.add_height(35)
    
    for i, m in enumerate(modelos, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(f"Modelo {i}  —  {m['titulo']}")
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = AZUL
        
        title_lines = 2 if len(m['titulo']) > 55 else 1
        tracker.add_height(17 + title_lines * 15)

        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(4)
        p2.paragraph_format.space_after = Pt(4)
        p2.paragraph_format.line_spacing = Pt(17)
        r2 = p2.add_run(m["descripcion"])
        r2.font.size = Pt(12)
        r2.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
        
        desc_lines = max(1, len(m["descripcion"]) // 65)
        tracker.add_height(8 + desc_lines * 17)

        p3 = doc.add_paragraph()
        p3.paragraph_format.space_before = Pt(4)
        p3.paragraph_format.space_after = Pt(6)
        r3a = p3.add_run("Recomendación:  ")
        r3a.bold = True
        r3a.font.size = Pt(12)
        r3a.font.color.rgb = AZUL_CLARO
        r3b = p3.add_run(m["recomendacion"])
        r3b.font.size = Pt(12)
        r3b.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
        
        rec_lines = max(1, len(m["recomendacion"]) // 65)
        tracker.add_height(10 + rec_lines * 17)

    # ── TENDENCIAS DEL MERCADO (Página 4–5) ──────────────────────────────────
    tracker.check_break_before_section(doc, section_height=130)
    add_section_header(doc, "📈 Tendencias del Mercado")
    tracker.add_height(35)
    
    for i, t in enumerate(tendencias, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(f"Tendencia {i}  —  {t['titulo']}")
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = AZUL
        
        title_lines = 2 if len(t['titulo']) > 55 else 1
        tracker.add_height(17 + title_lines * 15)

        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(4)
        p2.paragraph_format.space_after = Pt(8)
        p2.paragraph_format.line_spacing = Pt(17)
        r2 = p2.add_run(t["analisis"])
        r2.font.size = Pt(12)
        r2.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
        
        analysis_lines = max(1, len(t["analisis"]) // 65)
        tracker.add_height(12 + analysis_lines * 17)

    # ── VEREDICTO ACCIONABLE (Página 5) ──────────────────────────────────────
    tracker.check_break_before_section(doc, section_height=260)
    add_section_header(doc, "✅ Veredicto Accionable")
    tracker.add_height(35)
    
    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_before = Pt(6)
    p_intro.paragraph_format.space_after = Pt(6)
    r_intro = p_intro.add_run("5 acciones concretas para implementar esta semana:")
    r_intro.bold = True
    r_intro.font.size = Pt(12)
    r_intro.font.color.rgb = AZUL
    tracker.add_height(12 + 12)

    for i, item in enumerate(veredicto, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = Pt(17)
        p.paragraph_format.left_indent = Cm(0.5)
        r_num = p.add_run(f"{i}.  ")
        r_num.bold = True
        r_num.font.size = Pt(12)
        r_num.font.color.rgb = AZUL_CLARO
        r = p.add_run(item)
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
        
        item_lines = max(1, len(item) // 60)
        tracker.add_height(12 + item_lines * 17)

    # ── SUSCRIPCIÓN ───────────────────────────────────────────────────────────
    p_subs = doc.add_paragraph()
    p_subs.paragraph_format.space_before = Pt(14)
    p_subs.paragraph_format.space_after = Pt(6)
    r_s1 = p_subs.add_run("📩  Suscríbete gratis: ")
    r_s1.bold = True
    r_s1.font.size = Pt(12)
    r_s1.font.color.rgb = AZUL
    r_s2 = p_subs.add_run("www.emprendedores.ec/suscripcion")
    r_s2.font.size = Pt(12)
    r_s2.font.color.rgb = AZUL_CLARO
    r_s2.bold = True

    # ── HASHTAGS ──────────────────────────────────────────────────────────────
    add_section_header(doc, "# Comparte esta edición")
    todos = list(hashtags_fijos) + list(hashtags_dinamicos)
    p_ht = doc.add_paragraph()
    p_ht.paragraph_format.space_before = Pt(6)
    p_ht.paragraph_format.space_after = Pt(8)
    p_ht.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for i, tag in enumerate(todos):
        r = p_ht.add_run(tag)
        r.font.size = Pt(11)
        r.font.color.rgb = AZUL_CLARO
        r.bold = True
        if i < len(todos) - 1:
            sep = p_ht.add_run("  ")
            sep.font.size = Pt(11)

    # Cierre editorial
    p_cierre = doc.add_paragraph()
    p_cierre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cierre.paragraph_format.space_before = Pt(10)
    p_cierre.paragraph_format.space_after = Pt(4)
    r_c = p_cierre.add_run("— Hasta la próxima edición —")
    r_c.italic = True
    r_c.font.size = Pt(12)
    r_c.font.color.rgb = GRIS

    p_firma = doc.add_paragraph()
    p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_firma.paragraph_format.space_before = Pt(4)
    r_f = p_firma.add_run("Ricardo Wagner-Areco  ·  EMPRENDEDORES.LTD  ·  emprendedores.ec")
    r_f.font.size = Pt(11)
    r_f.bold = True
    r_f.font.color.rgb = AZUL

    # ── Guardar ───────────────────────────────────────────────────────────────
    out = Path(output_path)
    doc.save(out)
    patch_headers_xml(out, edition, date_str)
    print(f"✓ Generado: {out}")
    return out


def patch_headers_xml(docx_path, edition, date_str):
    """Reemplaza header/footer XML dentro del zip para garantizar formato exacto."""
    import zipfile, shutil, re

    NS = 'xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:w10="urn:schemas-microsoft-com:office:word" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml" mc:Ignorable="w14"'

    # ── XML del header vacío (página 1) ───────────────────────────────────────
    HDR_EMPTY = f"""<?xml version='1.0' encoding='UTF-8' standalone='yes'?>
<w:hdr {NS}><w:p><w:pPr><w:pStyle w:val="Header"/></w:pPr></w:p></w:hdr>"""

    # ── XML del header con contenido (páginas 2+) ─────────────────────────────
    # Diseñado con una tabla invisible de dos columnas para alinear perfectamente
    # el título y fecha a la izquierda y el nombre de la empresa a la derecha.
    HDR_CONTENT = f"""<?xml version='1.0' encoding='UTF-8' standalone='yes'?>
<w:hdr {NS}>
  <w:tbl>
    <w:tblPr>
      <w:tblW w:w="9406" w:type="dxa"/>
      <w:tblBorders>
        <w:top w:val="none"/>
        <w:left w:val="none"/>
        <w:bottom w:val="single" w:sz="6" w:space="4" w:color="003366"/>
        <w:right w:val="none"/>
        <w:insideH w:val="none"/>
        <w:insideV w:val="none"/>
      </w:tblBorders>
      <w:tblLayout w:type="fixed"/>
    </w:tblPr>
    <w:tblGrid>
      <w:gridCol w:w="6236"/>
      <w:gridCol w:w="3170"/>
    </w:tblGrid>
    <w:tr>
      <w:tc>
        <w:tcPr>
          <w:tcW w:w="6236" w:type="dxa"/>
        </w:tcPr>
        <w:p>
          <w:pPr>
            <w:pStyle w:val="Header"/>
            <w:jc w:val="left"/>
          </w:pPr>
          <w:r><w:rPr><w:b/><w:color w:val="003366"/><w:sz w:val="20"/></w:rPr><w:t xml:space="preserve">PULSO a la IA</w:t></w:r>
          <w:r><w:rPr><w:color w:val="606060"/><w:sz w:val="18"/></w:rPr><w:t xml:space="preserve">  ·  Edición {edition}  ·  {date_str}</w:t></w:r>
        </w:p>
      </w:tc>
      <w:tc>
        <w:tcPr>
          <w:tcW w:w="3170" w:type="dxa"/>
        </w:tcPr>
        <w:p>
          <w:pPr>
            <w:pStyle w:val="Header"/>
            <w:jc w:val="right"/>
          </w:pPr>
          <w:r><w:rPr><w:b/><w:color w:val="0070C0"/><w:sz w:val="20"/></w:rPr><w:t>EMPRENDEDORES.LTD</w:t></w:r>
        </w:p>
      </w:tc>
    </w:tr>
  </w:tbl>
</w:hdr>"""

    # ── XML del footer (mismo en página 1 y resto) ────────────────────────────
    FTR_CONTENT = f"""<?xml version='1.0' encoding='UTF-8' standalone='yes'?>
<w:ftr {NS}>
  <w:p>
    <w:pPr>
      <w:pStyle w:val="Footer"/>
      <w:jc w:val="left"/>
      <w:pBdr><w:top w:val="single" w:sz="4" w:space="4" w:color="003366"/></w:pBdr>
      <w:tabs>
        <w:tab w:val="center" w:pos="4536"/>
        <w:tab w:val="right" w:pos="9072"/>
      </w:tabs>
    </w:pPr>
    <w:r><w:rPr><w:color w:val="606060"/><w:sz w:val="16"/></w:rPr><w:t xml:space="preserve">emprendedores.ec  |  ricardowec@gmail.com</w:t></w:r>
    <w:r><w:tab/></w:r>
    <w:r><w:rPr><w:color w:val="606060"/><w:sz w:val="16"/></w:rPr><w:t xml:space="preserve">Pág. </w:t></w:r>
    <w:r><w:rPr><w:color w:val="606060"/><w:sz w:val="16"/></w:rPr><w:fldChar w:fldCharType="begin"/></w:r>
    <w:r><w:rPr><w:color w:val="606060"/><w:sz w:val="16"/></w:rPr><w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>
    <w:r><w:rPr><w:color w:val="606060"/><w:sz w:val="16"/></w:rPr><w:fldChar w:fldCharType="end"/></w:r>
    <w:r><w:tab/></w:r>
    <w:r><w:rPr><w:color w:val="0070C0"/><w:sz w:val="16"/></w:rPr><w:t>emprendedores.ec/suscripcion</w:t></w:r>
  </w:p>
</w:ftr>"""

    tmp = Path(str(docx_path) + ".tmp")
    with zipfile.ZipFile(docx_path, 'r') as zin, zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
        # Leer rels y document.xml para identificar qué archivo es "first" vs "default"
        rels_xml  = zin.read('word/_rels/document.xml.rels').decode('utf-8')
        doc_xml   = zin.read('word/document.xml').decode('utf-8')

        # Encontrar rId del header first y default
        first_rid   = re.search(r'headerReference[^/]*/>\s*' , doc_xml)  # fallback
        m_first   = re.search(r'headerReference\s+w:type="first"\s+r:id="([^"]+)"', doc_xml)
        m_default = re.search(r'headerReference\s+w:type="default"\s+r:id="([^"]+)"', doc_xml)
        m_fftr    = re.search(r'footerReference\s+w:type="first"\s+r:id="([^"]+)"', doc_xml)
        m_dftr    = re.search(r'footerReference\s+w:type="default"\s+r:id="([^"]+)"', doc_xml)

        def rid_to_file(rid):
            m = re.search(rf'Id="{rid}"[^>]+Target="([^"]+)"', rels_xml)
            return 'word/' + m.group(1) if m else None

        first_hdr_file   = rid_to_file(m_first.group(1))   if m_first   else None
        default_hdr_file = rid_to_file(m_default.group(1)) if m_default else None
        first_ftr_file   = rid_to_file(m_fftr.group(1))    if m_fftr    else None
        default_ftr_file = rid_to_file(m_dftr.group(1))    if m_dftr    else None

        replacements = {}
        if first_hdr_file:   replacements[first_hdr_file]   = HDR_EMPTY.encode('utf-8')
        if default_hdr_file: replacements[default_hdr_file] = HDR_CONTENT.encode('utf-8')
        if first_ftr_file:   replacements[first_ftr_file]   = FTR_CONTENT.encode('utf-8')
        if default_ftr_file: replacements[default_ftr_file] = FTR_CONTENT.encode('utf-8')

        for item in zin.infolist():
            data = replacements.get(item.filename)
            if data is None:
                data = zin.read(item.filename)
            zout.writestr(item, data)

    shutil.move(tmp, docx_path)


if __name__ == "__main__":
    import sys, json as _json
    if len(sys.argv) < 5:
        print("Uso: pulso_publisher.py <data.json> <salida.docx> <edition> <fecha>")
        sys.exit(1)
    _data = _json.loads(Path(sys.argv[1]).read_text())
    _resumen = _data.get("resumen_ejecutivo", "")
    _resumen_paras = [l.strip() for l in _resumen.split("\n") if l.strip()] if _resumen else [_resumen]
    generate(
        edition          = sys.argv[3],
        date_str         = sys.argv[4],
        resumen_ejecutivo= _resumen_paras,
        noticias         = _data.get("noticias", []),
        modelos          = _data.get("modelos", []),
        tendencias       = _data.get("tendencias", []),
        veredicto        = _data.get("veredicto", []),
        hashtags_fijos   = _data.get("hashtags_fijos", []),
        hashtags_dinamicos = _data.get("hashtags_dinamicos", []),
        output_path      = sys.argv[2],
    )
