#!/usr/bin/env python3
"""
PULSO a la IA — Curador de noticias
EMPRENDEDORES.LTD

Pipeline: RSS feeds → Claude Haiku (clasifica) → borrador .docx → formateador Node.js → Gmail
"""

import os, sys, json, yaml, hashlib, smtplib, logging, subprocess, re
from datetime import datetime, timedelta
from pathlib import Path
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders

import feedparser
import anthropic
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Paths relativos al script ──────────────────────────────────────────────────
BASE_DIR   = Path(__file__).resolve().parent
CONFIG     = BASE_DIR / "config.yaml"
CACHE_FILE = BASE_DIR / "cache" / "seen_articles.json"
LOG_FILE   = BASE_DIR / "logs" / "pulso.log"
OUTPUT_DIR = BASE_DIR / "output"

# ── Logging ────────────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[logging.FileHandler(LOG_FILE), logging.StreamHandler()]
)
log = logging.getLogger(__name__)


def load_config():
    with open(CONFIG) as f:
        return yaml.safe_load(f)


def load_cache():
    if CACHE_FILE.exists():
        return json.loads(CACHE_FILE.read_text())
    return {}


def save_cache(cache):
    CACHE_FILE.write_text(json.dumps(cache, indent=2))


def article_id(url):
    return hashlib.md5(url.encode()).hexdigest()


def fetch_feeds(cfg, cache):
    """Descarga y filtra artículos nuevos de todos los feeds."""
    lookback = datetime.now() - timedelta(days=cfg.get("lookback_days", 7))
    articles = []
    boost  = [k.lower() for k in cfg.get("boost_keywords", [])]
    suppress = [k.lower() for k in cfg.get("suppress_keywords", [])]

    for feed_cfg in cfg["feeds"]:
        url = feed_cfg["url"]
        try:
            parsed = feedparser.parse(url)
            for entry in parsed.entries:
                aid = article_id(entry.get("link", entry.get("title", "")))
                if aid in cache:
                    continue
                # Filtro de fecha
                pub = entry.get("published_parsed")
                if pub:
                    pub_dt = datetime(*pub[:6])
                    if pub_dt < lookback:
                        continue
                title   = entry.get("title", "")
                summary = entry.get("summary", entry.get("description", ""))
                text    = (title + " " + summary).lower()
                # Boost / suppress
                score = feed_cfg.get("priority", 2)
                for k in boost:
                    if k in text:
                        score = max(1, score - 1)
                for k in suppress:
                    if k in text:
                        score = 99  # efectivamente excluye
                if score == 99:
                    continue
                articles.append({
                    "id":       aid,
                    "title":    title,
                    "summary":  summary[:600],
                    "url":      entry.get("link", ""),
                    "source":   feed_cfg["name"],
                    "category": feed_cfg.get("category", "tech"),
                    "score":    score,
                })
        except Exception as e:
            log.warning(f"Error en feed {url}: {e}")

    articles.sort(key=lambda x: x["score"])
    return articles[:cfg.get("max_articles", 60)]


def classify_with_haiku(articles, cfg, edition_num):
    """Claude Haiku clasifica y redacta el borrador editorial."""
    client = anthropic.Anthropic(api_key=cfg["anthropic_api_key"])

    articles_text = "\n\n".join([
        f"[{i+1}] FUENTE: {a['source']} | CATEGORÍA: {a['category']}\n"
        f"TÍTULO: {a['title']}\nRESUMEN: {a['summary']}\nURL: {a['url']}"
        for i, a in enumerate(articles[:40])
    ])

    today = datetime.now().strftime("%d de %B de %Y")

    prompt = f"""Eres el editor de "PULSO a la IA", publicación semanal de EMPRENDEDORES.LTD para ejecutivos del sector bancario, comercial y agroindustrial en Ecuador y Latinoamérica.

Edición #{edition_num} — {today}

ARTÍCULOS DISPONIBLES:
{articles_text}

Tu tarea: Selecciona los artículos MÁS RELEVANTES para ejecutivos latinoamericanos y redacta el borrador completo de PULSO a la IA en español.

ESTRUCTURA OBLIGATORIA (respeta exactamente estos marcadores):

===RESUMEN_EJECUTIVO===
[2-3 párrafos con la tesis editorial de la semana. Qué tendencia dominante hay. Impacto para el sector financiero/agroindustrial latinoamericano.]

===NOTICIAS_DESTACADAS===
---NOTICIA_1---
TÍTULO: [título conciso]
QUÉ PASÓ: [1-2 oraciones]
POR QUÉ IMPORTA: [1-2 oraciones]
TE AFECTA: [impacto concreto para ejecutivo latinoamericano]
RECOMENDACIÓN: [acción concreta]
URL: [url fuente]
---NOTICIA_2---
[igual estructura]
---NOTICIA_3---
[igual estructura]

===MODELOS_DESTACADOS===
---MODELO_1---
TÍTULO: [nombre del modelo/herramienta]
DESCRIPCIÓN: [2-3 oraciones]
RECOMENDACIÓN: [para qué usarlo]
---MODELO_2---
[igual]

===TENDENCIAS_MERCADO===
---TENDENCIA_1---
TÍTULO: [título]
ANÁLISIS: [2-3 oraciones de análisis ejecutivo]
---TENDENCIA_2---
[igual]

===VEREDICTO_ACCIONABLE===
1. [Acción concreta 1 — máx 2 oraciones]
2. [Acción concreta 2]
3. [Acción concreta 3]
4. [Acción concreta 4]
5. [Acción concreta 5]

===FIN===

CRITERIOS:
- Priorizar noticias sobre banca, fintech, automatización, regulación IA, modelos de lenguaje
- Lenguaje ejecutivo, no técnico
- Perspectiva latinoamericana cuando sea posible
- Máximo 3 noticias, 2 modelos, 2 tendencias
"""

    log.info("Clasificando con Claude Haiku...")
    msg = client.messages.create(
        model="claude-haiku-4-5",
        max_tokens=4000,
        messages=[{"role": "user", "content": prompt}]
    )
    return msg.content[0].text


def parse_draft(draft_text):
    """Parsea el texto del borrador en secciones estructuradas."""
    data = {
        "resumen_ejecutivo": "",
        "noticias": [],
        "modelos": [],
        "tendencias": [],
        "veredicto": []
    }

    def extract_block(text, start_marker, end_marker):
        pattern = rf"{re.escape(start_marker)}(.*?){re.escape(end_marker)}"
        m = re.search(pattern, text, re.DOTALL)
        return m.group(1).strip() if m else ""

    data["resumen_ejecutivo"] = extract_block(
        draft_text, "===RESUMEN_EJECUTIVO===", "===NOTICIAS_DESTACADAS==="
    )

    def clean(text):
        """Elimina asteriscos markdown y espacios sobrantes."""
        return re.sub(r'\*+', '', text).strip()

    def find_field(chunk, *names):
        """Busca campo por múltiples variantes, tolerando asteriscos markdown."""
        for name in names:
            pattern = rf'\**\s*{re.escape(name)}\s*\**\s*:?\s*\**\s*(.+?)(?=\n\s*\**\s*[A-ZÁÉÍÓÚÑ]{{2,}}|$)'
            m = re.search(pattern, chunk, re.DOTALL | re.IGNORECASE)
            if m:
                return clean(m.group(1))
        return ""

    # Noticias
    noticias_block = extract_block(draft_text, "===NOTICIAS_DESTACADAS===", "===MODELOS_DESTACADOS===")
    for chunk in re.split(r"---NOTICIA_\d+---", noticias_block):
        chunk = chunk.strip()
        if not chunk: continue
        n = {
            "titulo":          find_field(chunk, "TÍTULO", "TITULO"),
            "que_paso":        find_field(chunk, "QUÉ PASÓ", "QUE PASO"),
            "por_que_importa": find_field(chunk, "POR QUÉ IMPORTA", "POR QUE IMPORTA"),
            "te_afecta":       find_field(chunk, "TE AFECTA"),
            "recomendacion":   find_field(chunk, "RECOMENDACIÓN", "RECOMENDACION"),
            "url":             find_field(chunk, "URL"),
        }
        if n.get("titulo"): data["noticias"].append(n)

    # Modelos
    modelos_block = extract_block(draft_text, "===MODELOS_DESTACADOS===", "===TENDENCIAS_MERCADO===")
    for chunk in re.split(r"---MODELO_\d+---", modelos_block):
        chunk = chunk.strip()
        if not chunk: continue
        m = {
            "titulo":        find_field(chunk, "TÍTULO", "TITULO"),
            "descripcion":   find_field(chunk, "DESCRIPCIÓN", "DESCRIPCION"),
            "recomendacion": find_field(chunk, "RECOMENDACIÓN", "RECOMENDACION"),
        }
        if m.get("titulo"): data["modelos"].append(m)

    # Tendencias
    tendencias_block = extract_block(draft_text, "===TENDENCIAS_MERCADO===", "===VEREDICTO_ACCIONABLE===")
    for chunk in re.split(r"---TENDENCIA_\d+---", tendencias_block):
        chunk = chunk.strip()
        if not chunk: continue
        t = {
            "titulo":  find_field(chunk, "TÍTULO", "TITULO"),
            "analisis":find_field(chunk, "ANÁLISIS", "ANALISIS"),
        }
        if t.get("titulo"): data["tendencias"].append(t)

    # Veredicto
    veredicto_block = extract_block(draft_text, "===VEREDICTO_ACCIONABLE===", "===FIN===")
    for line in veredicto_block.strip().splitlines():
        m = re.match(r"^\d+\.\s+(.+)", line.strip())
        if m: data["veredicto"].append(clean(m.group(1).strip()))

    return data


def generate_draft_docx(data, edition_num, output_path):
    """Genera .docx borrador simple (será formateado por el publisher)."""
    doc = Document()
    today = datetime.now().strftime("%d de %B de %Y")

    # Título
    t = doc.add_heading(f"PULSO a la IA — Edición {edition_num}", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Fecha: {today}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    # Resumen ejecutivo
    doc.add_heading("Resumen Ejecutivo", 1)
    doc.add_paragraph(data["resumen_ejecutivo"])
    doc.add_paragraph("")

    # Noticias
    doc.add_heading("1. Noticias Destacadas", 1)
    for i, n in enumerate(data["noticias"], 1):
        doc.add_heading(f"1.{i}  {n.get('titulo','')}", 2)
        if n.get("que_paso"):
            p = doc.add_paragraph()
            p.add_run("Qué pasó: ").bold = True
            p.add_run(n["que_paso"])
        if n.get("por_que_importa"):
            p = doc.add_paragraph()
            p.add_run("Por qué importa: ").bold = True
            p.add_run(n["por_que_importa"])
        if n.get("te_afecta"):
            p = doc.add_paragraph()
            p.add_run("Te afecta: ").bold = True
            p.add_run(n["te_afecta"])
        if n.get("recomendacion"):
            p = doc.add_paragraph()
            p.add_run("Recomendación: ").bold = True
            p.add_run(n["recomendacion"])
        doc.add_paragraph("")

    # Modelos
    if data["modelos"]:
        doc.add_heading("2. Modelos Destacados", 1)
        for i, m in enumerate(data["modelos"], 1):
            doc.add_heading(f"2.{i}  {m.get('titulo','')}", 2)
            if m.get("descripcion"):
                doc.add_paragraph(m["descripcion"])
            if m.get("recomendacion"):
                p = doc.add_paragraph()
                p.add_run("Recomendación: ").bold = True
                p.add_run(m["recomendacion"])
            doc.add_paragraph("")

    # Tendencias
    if data["tendencias"]:
        doc.add_heading("3. Tendencias del Mercado", 1)
        for i, t in enumerate(data["tendencias"], 1):
            doc.add_heading(f"3.{i}  {t.get('titulo','')}", 2)
            doc.add_paragraph(t.get("analisis", t.get("analsis", "")))
            doc.add_paragraph("")

    # Veredicto
    if data["veredicto"]:
        doc.add_heading("Veredicto Accionable", 1)
        for item in data["veredicto"]:
            doc.add_paragraph(item, style="List Number")
        doc.add_paragraph("")

    # Link suscripción
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("📩 Suscríbete a PULSO a la IA: ").bold = True
    p.add_run("https://emprendedores.ec/suscripcion")

    doc.save(output_path)
    # Guardar también como JSON para el publisher Node.js
    json_path = str(output_path).replace('.docx', '_data.json')
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    log.info(f"Borrador guardado: {output_path}")


def run_publisher(draft_path, edition_num, cfg):
    """Ejecuta el formateador Node.js (pulso_publisher.js)."""
    publisher_script = BASE_DIR / "pulso_publisher.js"
    output_path = OUTPUT_DIR / f"PULSO_a_la_IA_Edicion_{edition_num}.docx"

    if not publisher_script.exists():
        log.warning("pulso_publisher.js no encontrado — se usará el borrador sin formatear")
        return draft_path

    node_bin = cfg.get("node_bin", "node")
    result = subprocess.run(
        [node_bin, str(publisher_script),
         str(draft_path), str(output_path), str(edition_num)],
        capture_output=True, text=True, cwd=str(BASE_DIR)
    )
    if result.returncode == 0:
        log.info(f"Publisher OK → {output_path}")
        return output_path
    else:
        log.error(f"Publisher error: {result.stderr}")
        return draft_path


def send_email(docx_path, edition_num, cfg, dry_run=False):
    """Envía el .docx por SMTP."""
    ec = cfg["email"]
    today = datetime.now().strftime("%d/%m/%Y")

    msg = MIMEMultipart()
    msg["From"]    = ec["from"]
    msg["To"]      = ec["to"]
    msg["Subject"] = f"PULSO a la IA — Edición {edition_num} | {today}"

    body = f"""Estimado/a,

Adjunto encontrará la Edición {edition_num} de PULSO a la IA, su resumen semanal de tendencias en Inteligencia Artificial para el sector ejecutivo.

📌 Esta edición incluye:
• Noticias destacadas de IA con impacto en banca y agroindustria
• Modelos y herramientas relevantes
• Tendencias del mercado
• Veredicto accionable con recomendaciones concretas

📩 Suscríbete: https://emprendedores.ec/suscripcion

—
Ricardo Wagner-Areco
EMPRENDEDORES.LTD | emprendedores.ec
"""
    msg.attach(MIMEText(body, "plain", "utf-8"))

    # Adjunto
    with open(docx_path, "rb") as f:
        part = MIMEBase("application", "octet-stream")
        part.set_payload(f.read())
    encoders.encode_base64(part)
    part.add_header("Content-Disposition", f'attachment; filename="{docx_path.name}"')
    msg.attach(part)

    if dry_run:
        log.info(f"[DRY-RUN] Email preparado para {ec['to']} — no se envía")
        return

    with smtplib.SMTP(ec["smtp_host"], ec["smtp_port"]) as server:
        server.starttls()
        server.login(ec["smtp_user"], ec["smtp_password"])
        server.sendmail(ec["from"], ec["to"], msg.as_string())
    log.info(f"Email enviado a {ec['to']}")


def get_next_edition(cfg):
    """Lee/incrementa el número de edición."""
    edition_file = BASE_DIR / "cache" / "edition.txt"
    if edition_file.exists():
        n = int(edition_file.read_text().strip()) + 1
    else:
        n = cfg.get("start_edition", 1)
    edition_file.write_text(str(n))
    return n


def main():
    dry_run = "--dry-run" in sys.argv
    if dry_run:
        log.info("=== MODO DRY-RUN (no se envía email) ===")

    cfg = load_config()
    cache = load_cache()

    # 1. Fetch feeds
    log.info("Descargando feeds RSS...")
    articles = fetch_feeds(cfg, cache)
    log.info(f"  {len(articles)} artículos nuevos encontrados")

    if not articles:
        log.info("Sin artículos nuevos. Abortando.")
        return

    # 2. Clasificar con Haiku
    edition_num = get_next_edition(cfg)
    draft_text  = classify_with_haiku(articles, cfg, edition_num)

    # 3. Parsear borrador
    data = parse_draft(draft_text)
    log.info(f"  {len(data['noticias'])} noticias | {len(data['modelos'])} modelos | {len(data['tendencias'])} tendencias")

    # 4. Guardar borrador .docx simple
    draft_path = OUTPUT_DIR / f"borrador_edicion_{edition_num}.docx"
    generate_draft_docx(data, edition_num, draft_path)

    # 5. Ejecutar publisher formateador
    final_path = run_publisher(draft_path, edition_num, cfg)

    # 6. Enviar por email
    send_email(final_path, edition_num, cfg, dry_run=dry_run)

    # 7. Actualizar cache
    for a in articles:
        cache[a["id"]] = datetime.now().isoformat()
    save_cache(cache)

    log.info(f"=== PULSO Edición {edition_num} completado ===")


if __name__ == "__main__":
    main()

