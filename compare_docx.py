#!/usr/bin/env python3
import sys
import zipfile
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm

def get_xml_files(docx_path):
    """Extrae las partes XML relevantes del archivo .docx para su comparación directa."""
    xml_data = {}
    try:
        with zipfile.ZipFile(docx_path, 'r') as z:
            for name in ['word/document.xml', 'word/_rels/document.xml.rels']:
                if name in z.namelist():
                    xml_data[name] = z.read(name).decode('utf-8')
            
            # Buscar archivos de cabecera y pie de página
            for name in z.namelist():
                if name.startswith('word/header') or name.startswith('word/footer'):
                    xml_data[name] = z.read(name).decode('utf-8')
    except Exception as e:
        print(f"Error leyendo XML de {docx_path}: {e}")
    return xml_data

def analyze_document_structure(path):
    """Analiza la estructura del documento usando python-docx."""
    doc = Document(path)
    info = {
        "margins": {},
        "paragraphs_count": len(doc.paragraphs),
        "tables_count": len(doc.tables),
        "headings": [],
        "table_cells": [],
        "list_items": [],
    }
    
    # 1. Márgenes
    section = doc.sections[0]
    info["margins"] = {
        "top": section.top_margin.cm if section.top_margin else None,
        "bottom": section.bottom_margin.cm if section.bottom_margin else None,
        "left": section.left_margin.cm if section.left_margin else None,
        "right": section.right_margin.cm if section.right_margin else None,
    }
    
    # 2. Encabezados (Headings)
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            info["headings"].append({
                "text": p.text,
                "style": p.style.name,
                "alignment": str(p.alignment)
            })
            
    # 3. Tablas e Información de Celdas (Fondos y contenido)
    for idx, table in enumerate(doc.tables):
        table_info = {
            "index": idx,
            "rows": len(table.rows),
            "cols": len(table.columns),
            "cells": []
        }
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                # Extraer propiedades XML de la celda (como color de fondo)
                tcPr = cell._tc.get_or_add_tcPr()
                shd = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                bg_color = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill') if shd is not None else "None"
                
                # Párrafos en la celda
                paragraphs_text = [p.text for p in cell.paragraphs if p.text.strip()]
                
                table_info["cells"].append({
                    "row": r_idx,
                    "col": c_idx,
                    "bg_color": bg_color,
                    "text_snippet": " | ".join(paragraphs_text)[:150]
                })
        info["table_cells"].append(table_info)
        
    return info

def main():
    base_dir = Path(__file__).resolve().parent
    ref_path = base_dir / "output" / "PULSO_a_la_IA_Edicion_23.docx"
    test_path = base_dir / "output" / "TEST_FINAL.docx"
    
    print("======================================================================")
    print("🔍 COMPARANDO: TEST_FINAL.docx vs PULSO_a_la_IA_Edicion_23.docx (Referencia)")
    print("======================================================================\n")
    
    if not ref_path.exists():
        print(f"❌ Error: No se encontró el archivo de referencia en {ref_path}")
        sys.exit(1)
    if not test_path.exists():
        print(f"❌ Error: No se encontró el archivo test en {test_path}")
        sys.exit(1)
        
    # Análisis estructural
    ref_struct = analyze_document_structure(ref_path)
    test_struct = analyze_document_structure(test_path)
    
    # 1. Comparar Márgenes
    print("--- 📏 MÁRGENES ---")
    margins_match = True
    for key in ["top", "bottom", "left", "right"]:
        ref_val = ref_struct["margins"].get(key)
        test_val = test_struct["margins"].get(key)
        # Comparar con tolerancia flotante
        if ref_val is not None and test_val is not None and abs(ref_val - test_val) < 0.01:
            print(f"  ✅ Margen {key}: {ref_val:.2f} cm matches")
        else:
            print(f"  ❌ Margen {key}: Referencia={ref_val} cm | Test={test_val} cm")
            margins_match = False
            
    # 2. Comparar Tablas (Portada y Resumen Ejecutivo)
    print("\n--- 📊 TABLAS Y RECUADROS ---")
    print(f"  Tablas en Referencia: {ref_struct['tables_count']} | en Test: {test_struct['tables_count']}")
    
    # Comparar fondo del banner de título (Tabla 0)
    if len(ref_struct["table_cells"]) > 0 and len(test_struct["table_cells"]) > 0:
        ref_bg = ref_struct["table_cells"][0]["cells"][0]["bg_color"]
        test_bg = test_struct["table_cells"][0]["cells"][0]["bg_color"]
        if ref_bg.upper() == test_bg.upper():
            print(f"  ✅ Banner Título (Tabla 1) Fondo: {ref_bg} matches")
        else:
            print(f"  ❌ Banner Título Fondo: Referencia={ref_bg} | Test={test_bg}")
            
    # Comparar fondo del Resumen Ejecutivo (Tabla 1)
    if len(ref_struct["table_cells"]) > 1 and len(test_struct["table_cells"]) > 1:
        ref_bg = ref_struct["table_cells"][1]["cells"][0]["bg_color"]
        test_bg = test_struct["table_cells"][1]["cells"][0]["bg_color"]
        if ref_bg.upper() == test_bg.upper():
            print(f"  ✅ Recuadro Resumen Ejecutivo (Tabla 2) Fondo: {ref_bg} matches")
        else:
            print(f"  ❌ Recuadro Resumen Ejecutivo Fondo: Referencia={ref_bg} | Test={test_bg}")
            
    # 3. Comparar XML interno (Headers, Footers y Rels)
    print("\n--- ✉️ CABECERAS Y PIES DE PÁGINA (XML) ---")
    ref_xml = get_xml_files(ref_path)
    test_xml = get_xml_files(test_path)
    
    # Comparar qué archivos XML existen en cada zip
    ref_files = set(ref_xml.keys())
    test_files = set(test_xml.keys())
    
    missing_in_test = ref_files - test_files
    extra_in_test = test_files - ref_files
    
    if not missing_in_test and not extra_in_test:
        print("  ✅ Los mismos archivos XML de cabecera y pie de página existen en ambos zip.")
    else:
        if missing_in_test:
            print(f"  ❌ Faltan en el Test: {missing_in_test}")
        if extra_in_test:
            print(f"  ❌ Extras en el Test: {extra_in_test}")
            
    # Verificar si el header principal tiene la marca 'PULSO a la IA' y 'EMPRENDEDORES.LTD'
    print("\n--- 🔍 VERIFICACIONES ESPECÍFICAS DE FORMATO ---")
    
    # Comprobar cabecera principal (debería ser header2.xml normalmente)
    header_content_checked = False
    for filename, content in test_xml.items():
        if 'header' in filename:
            if 'PULSO a la IA' in content and 'EMPRENDEDORES.LTD' in content:
                print(f"  ✅ Cabecera dinámica encontrada en `{filename}` con textos de marca correctos.")
                header_content_checked = True
                
    if not header_content_checked:
        print("  ❌ No se encontró una cabecera con el texto corporativo esperado ('PULSO a la IA' y 'EMPRENDEDORES.LTD').")
        
    # Comprobar pie de página principal
    footer_content_checked = False
    for filename, content in test_xml.items():
        if 'footer' in filename:
            if 'emprendedores.ec' in content and 'PAGE' in content:
                print(f"  ✅ Pie de página dinámico encontrado en `{filename}` con 'emprendedores.ec' y campo PAGE.")
                footer_content_checked = True
                
    if not footer_content_checked:
        print("  ❌ No se encontró un pie de página con el texto o número de página dinámico (PAGE) esperado.")
        
    # 4. Verificar ausencia de tablas en la sección Veredicto
    print("\n--- 🕵️ VEREDICTO ACCIONABLE Y HASHTAGS ---")
    doc_test = Document(test_path)
    veredicto_found = False
    veredicto_has_table = False
    for p in doc_test.paragraphs:
        if "veredicto" in p.text.lower() or "acción" in p.text.lower():
            veredicto_found = True
            
    # El veredicto no debería estar dentro de una tabla en TEST_FINAL (sección libre con lista numerada)
    # Comprobamos si las tablas contienen palabras típicas del veredicto
    for idx, table in enumerate(doc_test.tables):
        if idx > 1: # Las primeras dos tablas son Portada y Resumen Ejecutivo
            for row in table.rows:
                for cell in row.cells:
                    if "acciones" in cell.text.lower() or "evalúa" in cell.text.lower():
                        veredicto_has_table = True
                        
    if veredicto_found:
        if not veredicto_has_table:
            print("  ✅ Veredicto Accionable formateado correctamente como lista directa (sin recuadro/tabla adicional).")
        else:
            print("  ❌ Alerta: Se detectó que el Veredicto Accionable está contenido en una tabla.")
    else:
        print("  ⚠️ No se localizó claramente la sección del Veredicto en el documento test.")

if __name__ == "__main__":
    main()
