#!/usr/bin/env python3
"""
Generador de la Edición Especial: ANTIGRAVITY 2.0: Mi experiencia
27 de mayo de 2026 | EMPRENDEDORES.LTD
Genera un documento Word corporativo único con la identidad visual de PULSO a la IA.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = BASE_DIR / "output"
OUTPUT_DIR.mkdir(exist_ok=True)

# ── Configuración Visual ───────────────────────────────────────────────────────
AZUL = RGBColor(0x00, 0x33, 0x66)        # Azul EMPRENDEDORES.LTD
AZUL_CLARO = RGBColor(0x00, 0x70, 0xC0)  # Azul secundario
GRIS = RGBColor(0x60, 0x60, 0x60)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)

# ── Datos de la Edición Especial ───────────────────────────────────────────────
EDITION_LABEL = "Edición Especial"
DATE_STR = "27 de mayo de 2026"
DOC_TITLE = "ANTIGRAVITY 2.0: Mi experiencia"
DOC_SUBTITLE = "Transición, resiliencia y el poder de la IA agéntica local"

RESUMEN_EJECUTIVO = [
    "La transición hacia flujos de desarrollo asistidos por Inteligencia Artificial (IA) ha demostrado ser una de las tendencias más disruptivas en la productividad ejecutiva y técnica de 2026. Sin embargo, la dependencia absoluta de servicios en la nube expone a los proyectos a riesgos operativos críticos, tales como los límites repentinos de cuotas de API y la volatilidad del contexto, los cuales pueden interrumpir el desarrollo en momentos clave de la integración.",
    "Frente a estas limitaciones, la arquitectura agéntica de Antigravity 2.0, impulsada por Gemini 3.5, introduce un paradigma de resiliencia y adaptabilidad en la Mac Mini M4 Pro. Al sincronizar de forma nativa el contexto local, no solo estabiliza el código tras cortes inesperados, sino que eleva el valor del proyecto mediante la creación de dashboards de control interactivos que garantizan la supervisión y continuidad del negocio."
]

HITOS = [
    {
        "titulo": "El arranque con Claude-Code y la transición de contexto",
        "descripcion": "El proyecto se inició con éxito bajo Claude-Code con el modelo Sonnet 4.7. Tras una sesión de desarrollo intensa y productiva, las cuotas límites de API de la cuenta detuvieron la ejecución. Al intentar reanudar, las diferencias de contexto heredadas dificultaron la estabilización inmediata de la sesión, un reto común al alternar de forma activa el desarrollo asistido.",
        "aprendizaje": "Las herramientas basadas en límites de cuotas de API web requieren planificar la persistencia de estado para evitar la fricción al alternar sesiones de desarrollo."
    },
    {
        "titulo": "El relevo resiliente de Antigravity 2.0 y Gemini 3.5",
        "descripcion": "Ante la inestabilidad del entorno anterior, se realizó la transición a Antigravity 2.0 equipado con el motor Gemini 3.5. La nueva plataforma demostró una capacidad excepcional de recuperación de contexto local, sincronizando el estado del repositorio y retomando el desarrollo desde un punto estable sin requerir explicaciones redundantes o reentrenamiento manual del agente.",
        "aprendizaje": "La resiliencia en la gestión de contexto y la adaptabilidad ante fallos de herramientas previas son determinantes para el éxito del desarrollo de software asistido por IA."
    },
    {
        "titulo": "Creación del Dashboard de Control y Automatización",
        "descripcion": "De manera proactiva, Antigravity 2.0 identificó la oportunidad de mejorar la experiencia del usuario diseñando e implementando un Dashboard de control web local (React + Express). Esta herramienta permite monitorear y disparar manualmente las tareas de curación con Gemini y compilación en Word, culminando con un botón estilizado de apagado del servidor para optimizar el consumo de recursos de la Mac Mini M4 Pro.",
        "aprendizaje": "Las soluciones de IA no deben ser solo scripts aislados, sino plataformas integradas que empoderen al usuario a supervisar de forma visual e intuitiva el pipeline autónomo."
    }
]

VEREDICTO = [
    "Implementa un sistema de logging persistente y registro de actividades en archivos markdown (ej. *.md) en tus proyectos para permitir a cualquier agente retomar la tarea sin pérdida de contexto.",
    "Diversifica los proveedores de modelos de lenguaje (LLM) en tus pipelines agénticos para asegurar redundancia en caso de caídas de servicio o límites de cuotas de API.",
    "Aprovecha la capacidad de cómputo local (como los 64 GB de la Mac Mini M4 Pro) para desplegar arquitecturas híbridas de desarrollo, reduciendo la dependencia de latencias en la nube.",
    "Diseña interfaces gráficas simples (Dashboards locales) para tus pipelines de datos complejos; la visibilidad del estado reduce la fricción y acelera la toma de decisiones ejecutivas.",
    "Define siempre planes de trabajo claros y solicita autorización previa antes de la ejecución de tareas críticas en entornos autónomos para mantener el control editorial del producto final."
]

HASHTAGS_FIJOS = [
    "#PulsoALaIA", "#InteligenciaArtificial", "#IA",
    "#EMPRENDEDORES", "#Newsletter", "#Tecnologia", "#LiderazgoDigital"
]

HASHTAGS_DINAMICOS = [
    "#Antigravity20", "#Gemini35", "#MacMiniM4Pro",
    "#BancaDigital", "#DesarrolloAgentico", "#ResilienciaTI"
]


# ── Utilidades de formato ──────────────────────────────────────────────────────

def truncate_text(text, max_len=55):
    if not text:
        return ""
    if len(text) <= max_len:
        return text
    truncated = text[:max_len]
    last_space = truncated.rfind(' ')
    if last_space > 0:
        truncated = truncated[:last_space]
    return truncated.strip() + "..."


def format_list_index(titles, max_len=50):
    if not titles:
        return ""
    if len(titles) == 1:
        return truncate_text(titles[0], max_len)
    elif len(titles) == 2:
        joined = f"{titles[0]} y {titles[1]}"
        if len(joined) <= max_len:
            return joined
    first = truncate_text(titles[0], max_len - 12)
    return f"{first} (y {len(titles)-1} más)"


class LayoutTracker:
    def __init__(self, page_height_limit=620):  # Alto útil real de página carta (puntos)
        self.current_height = 0
        self.limit = 310  # Límite del 50% de la página (la mitad)

    def add_height(self, points):
        if self.current_height + points <= 620:
            self.current_height += points
        else:
            # Word empuja el bloque entero a la siguiente página.
            # Por lo tanto, la altura en la nueva página empieza exactamente con el tamaño de este bloque.
            self.current_height = points

    def force_page_break(self):
        self.current_height = 0

    def check_break_before_section(self, doc, section_height=130):
        # Desactivado a petición del usuario para dejar que Word maneje el flujo libremente.
        return False


def add_section_header(doc, text, space_before=14, space_after=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = AZUL
    border = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '003366')
    border.append(bottom)
    p._p.get_or_add_pPr().append(border)
    return p


def set_margins(doc, top=2.5, bottom=2.2, left=2.5, right=2.5):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)


def add_page_number(para):
    r_begin = OxmlElement('w:r')
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    r_begin.append(fld_begin)
    para._p.append(r_begin)

    r_instr = OxmlElement('w:r')
    instr = OxmlElement('w:instrText')
    instr.text = ' PAGE '
    r_instr.append(instr)
    para._p.append(r_instr)

    r_end = OxmlElement('w:r')
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r_end.append(fld_end)
    para._p.append(r_end)


def build_footer_para(para):
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = para._p.get_or_add_pPr()

    pBdr = OxmlElement('w:pBdr')
    top_b = OxmlElement('w:top')
    top_b.set(qn('w:val'), 'single')
    top_b.set(qn('w:sz'), '4')
    top_b.set(qn('w:space'), '4')
    top_b.set(qn('w:color'), '003366')
    pBdr.append(top_b)
    pPr.append(pBdr)

    tabs = OxmlElement('w:tabs')
    tc = OxmlElement('w:tab')
    tc.set(qn('w:val'), 'center')
    tc.set(qn('w:pos'), '4536')
    tabs.append(tc)
    tr = OxmlElement('w:tab')
    tr.set(qn('w:val'), 'right')
    tr.set(qn('w:pos'), '9072')
    tabs.append(tr)
    pPr.append(tabs)

    r1 = para.add_run("emprendedores.ec  |  ricardowec@gmail.com")
    r1.font.size = Pt(8)
    r1.font.color.rgb = GRIS

    tab1 = OxmlElement('w:r')
    tab1.append(OxmlElement('w:tab'))
    para._p.append(tab1)
    r_pn = para.add_run("Pág. ")
    r_pn.font.size = Pt(8)
    r_pn.font.color.rgb = GRIS
    add_page_number(para)

    tab2 = OxmlElement('w:r')
    tab2.append(OxmlElement('w:tab'))
    para._p.append(tab2)
    r_fr = para.add_run("emprendedores.ec/suscripcion")
    r_fr.font.size = Pt(8)
    r_fr.font.color.rgb = AZUL_CLARO


def _remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBdr = OxmlElement('w:tcBdr')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tcBdr.append(b)
    tcPr.append(tcBdr)


def _add_para_bottom_border(para, color='003366'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '4')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def _set_cell_bg(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def _set_cell_margins(cell, top=280, left=280, bottom=280, right=280):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)


def _make_no_border_table(doc, rows=1, cols=1):
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl_xml = tbl._tbl
    tblPr = tbl_xml.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl_xml.insert(0, tblPr)

    tblGrid = tbl_xml.find(qn('w:tblGrid'))
    if tblGrid is not None:
        total_w = sum(int(gc.get(qn('w:w'), 0)) for gc in tblGrid.findall(qn('w:gridCol')))
    else:
        s = doc.sections[0]
        total_w = int((s.page_width - s.left_margin - s.right_margin) / 635)

    existing = tblPr.find(qn('w:tblW'))
    if existing is not None:
        tblPr.remove(existing)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(total_w))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    tblBdr = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        tblBdr.append(b)
    tblPr.append(tblBdr)
    return tbl


def add_title_block(doc, title, subtitle, label, date_str):
    tbl = _make_no_border_table(doc)
    cell = tbl.cell(0, 0)
    _remove_cell_borders(cell)
    _set_cell_bg(cell, '003366')
    _set_cell_margins(cell, top=300, left=440, bottom=300, right=440)

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(f"PULSO a la IA  ·  {label}")
    r1.bold = True
    r1.font.size = Pt(24)
    r1.font.color.rgb = BLANCO

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(3)
    r2 = p2.add_run(f"{title}\n{subtitle}")
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0xCC, 0xE5, 0xFF)

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(2)
    p3.paragraph_format.space_after = Pt(4)
    r3 = p3.add_run(f"EMPRENDEDORES.LTD  ·  {date_str}")
    r3.bold = True
    r3.font.size = Pt(11)
    r3.font.color.rgb = RGBColor(0x66, 0xB8, 0xFF)


def add_summary_block(doc, paragraphs_text):
    tbl = _make_no_border_table(doc)
    cell = tbl.cell(0, 0)
    _remove_cell_borders(cell)
    _set_cell_bg(cell, 'E3F2FD')
    _set_cell_margins(cell, top=180, left=360, bottom=180, right=360)

    p_title = cell.paragraphs[0]
    p_title.paragraph_format.space_before = Pt(1)
    p_title.paragraph_format.space_after = Pt(3)
    r_t = p_title.add_run("📊 RESUMEN EJECUTIVO")
    r_t.bold = True
    r_t.font.size = Pt(11)
    r_t.font.color.rgb = AZUL
    _add_para_bottom_border(p_title)

    for texto in paragraphs_text:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = Pt(18)
        r = p.add_run(texto)
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)


def _build_header_table(container, label, date_str):
    for p in container.paragraphs:
        p.clear()

    tbl = container.add_table(rows=1, cols=2, width=Cm(16))
    tbl_xml = tbl._tbl
    tblPr = tbl_xml.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl_xml.insert(0, tblPr)
    tblBdr = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        tblBdr.append(b)
    tblPr.append(tblBdr)

    for i, width_cm in enumerate([11, 5]):
        cell = tbl.cell(0, i)
        _remove_cell_borders(cell)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(int(width_cm * 567)))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

    p_l = tbl.cell(0, 0).paragraphs[0]
    pPr_l = p_l._p.get_or_add_pPr()
    jc_l = OxmlElement('w:jc')
    jc_l.set(qn('w:val'), 'left')
    pPr_l.append(jc_l)
    r1 = p_l.add_run("PULSO a la IA")
    r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = AZUL
    r2 = p_l.add_run(f"  ·  {label}  ·  {date_str}")
    r2.font.size = Pt(9); r2.font.color.rgb = GRIS

    p_r = tbl.cell(0, 1).paragraphs[0]
    pPr_r = p_r._p.get_or_add_pPr()
    jc_r = OxmlElement('w:jc')
    jc_r.set(qn('w:val'), 'right')
    pPr_r.append(jc_r)
    r3 = p_r.add_run("EMPRENDEDORES.LTD")
    r3.bold = True; r3.font.size = Pt(10); r3.font.color.rgb = AZUL_CLARO

    last_p = container.paragraphs[-1]
    _add_para_bottom_border(last_p)


def setup_header_footer(doc, label, date_str):
    section = doc.sections[0]
    sectPr = section._sectPr
    if sectPr.find(qn('w:titlePg')) is None:
        titlePg = OxmlElement('w:titlePg')
        sectPr.append(titlePg)

    first_hdr = section.first_page_header
    first_hdr.is_linked_to_previous = False
    for p in first_hdr.paragraphs:
        p.clear()

    first_ftr = section.first_page_footer
    first_ftr.is_linked_to_previous = False
    for p in first_ftr.paragraphs:
        p.clear()
    build_footer_para(first_ftr.paragraphs[0])

    _build_header_table(section.header, label, date_str)

    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()
    build_footer_para(footer.paragraphs[0])


def patch_headers_xml(docx_path, label, date_str):
    import zipfile, shutil, re

    NS = 'xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:w10="urn:schemas-microsoft-com:office:word" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml" mc:Ignorable="w14"'

    HDR_EMPTY = f"""<?xml version='1.0' encoding='UTF-8' standalone='yes'?>
<w:hdr {NS}><w:p><w:pPr><w:pStyle w:val="Header"/></w:pPr></w:p></w:hdr>"""

    # ── XML del header con contenido (páginas 2+) ─────────────────────────────
    # Diseñado con una tabla invisible de dos columnas para alinear perfectamente
    # el título y fecha a la izquierda y el nombre de la empresa a la derecha.
    HDR_CONTENT = f"""<?xml version='1.0' encoding='UTF-8' standalone='yes'?>
<w:hdr {NS}>
  <w:tbl>
    <w:tblPr>
      <w:tblW w:w="9406" w:type="dxa"/>
      <w:tblBorders>
        <w:top w:val="none"/>
        <w:left w:val="none"/>
        <w:bottom w:val="single" w:sz="6" w:space="4" w:color="003366"/>
        <w:right w:val="none"/>
        <w:insideH w:val="none"/>
        <w:insideV w:val="none"/>
      </w:tblBorders>
      <w:tblLayout w:type="fixed"/>
    </w:tblPr>
    <w:tblGrid>
      <w:gridCol w:w="6236"/>
      <w:gridCol w:w="3170"/>
    </w:tblGrid>
    <w:tr>
      <w:tc>
        <w:tcPr>
          <w:tcW w:w="6236" w:type="dxa"/>
        </w:tcPr>
        <w:p>
          <w:pPr>
            <w:pStyle w:val="Header"/>
            <w:jc w:val="left"/>
          </w:pPr>
          <w:r><w:rPr><w:b/><w:color w:val="003366"/><w:sz w:val="20"/></w:rPr><w:t xml:space="preserve">PULSO a la IA</w:t></w:r>
          <w:r><w:rPr><w:color w:val="606060"/><w:sz w:val="18"/></w:rPr><w:t xml:space="preserve">  ·  {label}  ·  {date_str}</w:t></w:r>
        </w:p>
      </w:tc>
      <w:tc>
        <w:tcPr>
          <w:tcW w:w="3170" w:type="dxa"/>
        </w:tcPr>
        <w:p>
          <w:pPr>
            <w:pStyle w:val="Header"/>
            <w:jc w:val="right"/>
          </w:pPr>
          <w:r><w:rPr><w:b/><w:color w:val="0070C0"/><w:sz w:val="20"/></w:rPr><w:t>EMPRENDEDORES.LTD</w:t></w:r>
        </w:p>
      </w:tc>
    </w:tr>
  </w:tbl>
</w:hdr>"""

    FTR_CONTENT = f"""<?xml version='1.0' encoding='UTF-8' standalone='yes'?>
<w:ftr {NS}>
  <w:p>
    <w:pPr>
      <w:pStyle w:val="Footer"/>
      <w:jc w:val="left"/>
      <w:pBdr><w:top w:val="single" w:sz="4" w:space="4" w:color="003366"/></w:pBdr>
      <w:tabs>
        <w:tab w:val="center" w:pos="4536"/>
        <w:tab w:val="right" w:pos="9072"/>
      </w:tabs>
    </w:pPr>
    <w:r><w:rPr><w:color w:val="606060"/><w:sz w:val="16"/></w:rPr><w:t xml:space="preserve">emprendedores.ec  |  ricardowec@gmail.com</w:t></w:r>
    <w:r><w:tab/></w:r>
    <w:r><w:rPr><w:color w:val="606060"/><w:sz w:val="16"/></w:rPr><w:t xml:space="preserve">Pág. </w:t></w:r>
    <w:r><w:rPr><w:color w:val="606060"/><w:sz w:val="16"/></w:rPr><w:fldChar w:fldCharType="begin"/></w:r>
    <w:r><w:rPr><w:color w:val="606060"/><w:sz w:val="16"/></w:rPr><w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>
    <w:r><w:rPr><w:color w:val="606060"/><w:sz w:val="16"/></w:rPr><w:fldChar w:fldCharType="end"/></w:r>
    <w:r><w:tab/></w:r>
    <w:r><w:rPr><w:color w:val="0070C0"/><w:sz w:val="16"/></w:rPr><w:t>emprendedores.ec/suscripcion</w:t></w:r>
  </w:p>
</w:ftr>"""

    tmp = Path(str(docx_path) + ".tmp")
    with zipfile.ZipFile(docx_path, 'r') as zin, zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
        rels_xml = zin.read('word/_rels/document.xml.rels').decode('utf-8')
        doc_xml = zin.read('word/document.xml').decode('utf-8')

        m_first = re.search(r'headerReference\s+w:type="first"\s+r:id="([^"]+)"', doc_xml)
        m_default = re.search(r'headerReference\s+w:type="default"\s+r:id="([^"]+)"', doc_xml)
        m_fftr = re.search(r'footerReference\s+w:type="first"\s+r:id="([^"]+)"', doc_xml)
        m_dftr = re.search(r'footerReference\s+w:type="default"\s+r:id="([^"]+)"', doc_xml)

        def rid_to_file(rid):
            m = re.search(rf'Id="{rid}"[^>]+Target="([^"]+)"', rels_xml)
            return 'word/' + m.group(1) if m else None

        first_hdr_file = rid_to_file(m_first.group(1)) if m_first else None
        default_hdr_file = rid_to_file(m_default.group(1)) if m_default else None
        first_ftr_file = rid_to_file(m_fftr.group(1)) if m_fftr else None
        default_ftr_file = rid_to_file(m_dftr.group(1)) if m_dftr else None

        replacements = {}
        if first_hdr_file:   replacements[first_hdr_file] = HDR_EMPTY.encode('utf-8')
        if default_hdr_file: replacements[default_hdr_file] = HDR_CONTENT.encode('utf-8')
        if first_ftr_file:   replacements[first_ftr_file] = FTR_CONTENT.encode('utf-8')
        if default_ftr_file: replacements[default_ftr_file] = FTR_CONTENT.encode('utf-8')

        for item in zin.infolist():
            data = replacements.get(item.filename)
            if data is None:
                data = zin.read(item.filename)
            zout.writestr(item, data)

    shutil.move(tmp, docx_path)


def generate_document(output_path):
    doc = Document()
    set_margins(doc)
    setup_header_footer(doc, EDITION_LABEL, DATE_STR)

    if len(doc.paragraphs) > 0 and doc.paragraphs[0].text == "":
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    doc.styles['Normal'].font.size = Pt(12)
    styles_el = doc.styles.element
    doc_defaults = styles_el.find(qn('w:docDefaults'))
    if doc_defaults is not None:
        rPr = doc_defaults.find('.//' + qn('w:rPr'))
        if rPr is not None:
            for tag in [qn('w:sz'), qn('w:szCs')]:
                el = rPr.find(tag)
                if el is not None:
                    el.set(qn('w:val'), '24')  # 12pt

    # ── PORTADA ───────────────────────────────────────────────────────────────
    add_title_block(doc, DOC_TITLE, DOC_SUBTITLE, EDITION_LABEL, DATE_STR)

    # Índice de contenidos
    add_section_header(doc, "En esta edición", space_before=8, space_after=2)
    
    resumen_desc = f"Tesis: {truncate_text(RESUMEN_EJECUTIVO[0], 50)}" if RESUMEN_EJECUTIVO else "Tesis de la edición"
    hitos_desc = format_list_index([h['titulo'] for h in HITOS], 50)
    veredicto_desc = f"{len(VEREDICTO)} conclusiones y aprendizajes clave"
    
    indice = [
        ("📊", "Resumen Ejecutivo", resumen_desc),
        ("🛠️", "Hitos de Desarrollo", hitos_desc),
        ("✅", "Veredicto Accionable", veredicto_desc),
    ]
    for emoji, titulo, desc in indice:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r1 = p.add_run(f"{emoji}  {titulo}: ")
        r1.bold = True
        r1.font.size = Pt(12)
        r1.font.color.rgb = AZUL
        r2 = p.add_run(desc)
        r2.font.size = Pt(12)
        r2.font.color.rgb = GRIS

    # ── RESUMEN EJECUTIVO ─────────────────────────────────────────────────────
    add_summary_block(doc, RESUMEN_EJECUTIVO)

    # Inicializar tracker de diseño para las siguientes páginas
    tracker = LayoutTracker()

    # ── HITOS DE DESARROLLO (Páginas 2+) ─────────────────────────────────────
    doc.add_page_break()
    tracker.force_page_break()
    add_section_header(doc, "🛠️ Hitos de Desarrollo")
    tracker.add_height(35)
    
    for i, h in enumerate(HITOS, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(f"Hito {i}  —  {h['titulo']}")
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = AZUL
        
        title_lines = 2 if len(h['titulo']) > 55 else 1
        tracker.add_height(17 + title_lines * 15)

        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(4)
        p2.paragraph_format.space_after = Pt(4)
        p2.paragraph_format.line_spacing = Pt(17)
        r2 = p2.add_run(h["descripcion"])
        r2.font.size = Pt(12)
        r2.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
        
        desc_lines = max(1, len(h["descripcion"]) // 65)
        tracker.add_height(8 + desc_lines * 17)

        p3 = doc.add_paragraph()
        p3.paragraph_format.space_before = Pt(4)
        p3.paragraph_format.space_after = Pt(6)
        r3a = p3.add_run("Aprendizaje:  ")
        r3a.bold = True
        r3a.font.size = Pt(12)
        r3a.font.color.rgb = AZUL_CLARO
        r3b = p3.add_run(h["aprendizaje"])
        r3b.font.size = Pt(12)
        r3b.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
        
        rec_lines = max(1, len(h["aprendizaje"]) // 65)
        tracker.add_height(10 + rec_lines * 17)

    # ── VEREDICTO ACCIONABLE ──────────────────────────────────────────────────
    tracker.check_break_before_section(doc, section_height=260)
    add_section_header(doc, "✅ Veredicto Accionable")
    tracker.add_height(35)
    
    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_before = Pt(6)
    p_intro.paragraph_format.space_after = Pt(6)
    r_intro = p_intro.add_run("Recomendaciones clave basadas en esta experiencia:")
    r_intro.bold = True
    r_intro.font.size = Pt(12)
    r_intro.font.color.rgb = AZUL
    tracker.add_height(12 + 12)

    for i, item in enumerate(VEREDICTO, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = Pt(17)
        p.paragraph_format.left_indent = Cm(0.5)
        r_num = p.add_run(f"{i}.  ")
        r_num.bold = True
        r_num.font.size = Pt(12)
        r_num.font.color.rgb = AZUL_CLARO
        r = p.add_run(item)
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
        
        item_lines = max(1, len(item) // 60)
        tracker.add_height(12 + item_lines * 17)

    # ── SUSCRIPCIÓN ───────────────────────────────────────────────────────────
    p_subs = doc.add_paragraph()
    p_subs.paragraph_format.space_before = Pt(14)
    p_subs.paragraph_format.space_after = Pt(6)
    r_s1 = p_subs.add_run("📩  Suscríbete gratis: ")
    r_s1.bold = True; r_s1.font.size = Pt(12); r_s1.font.color.rgb = AZUL
    r_s2 = p_subs.add_run("www.emprendedores.ec/suscripcion")
    r_s2.font.size = Pt(12); r_s2.font.color.rgb = AZUL_CLARO; r_s2.bold = True

    # ── HASHTAGS ──────────────────────────────────────────────────────────────
    add_section_header(doc, "# Comparte esta edición")
    todos = list(HASHTAGS_FIJOS) + list(HASHTAGS_DINAMICOS)
    p_ht = doc.add_paragraph()
    p_ht.paragraph_format.space_before = Pt(6)
    p_ht.paragraph_format.space_after = Pt(8)
    p_ht.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for idx, tag in enumerate(todos):
        r = p_ht.add_run(tag)
        r.font.size = Pt(11); r.font.color.rgb = AZUL_CLARO; r.bold = True
        if idx < len(todos) - 1:
            sep = p_ht.add_run("  ")
            sep.font.size = Pt(11)

    p_cierre = doc.add_paragraph()
    p_cierre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cierre.paragraph_format.space_before = Pt(10)
    p_cierre.paragraph_format.space_after = Pt(4)
    r_c = p_cierre.add_run("— Edición Especial —")
    r_c.italic = True; r_c.font.size = Pt(12); r_c.font.color.rgb = GRIS

    p_firma = doc.add_paragraph()
    p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_firma.paragraph_format.space_before = Pt(4)
    r_f = p_firma.add_run("Ricardo Wagner-Areco  ·  EMPRENDEDORES.LTD  ·  emprendedores.ec")
    r_f.font.size = Pt(11); r_f.bold = True; r_f.font.color.rgb = AZUL

    # Guardar
    out_path = Path(output_path)
    doc.save(out_path)
    patch_headers_xml(out_path, EDITION_LABEL, DATE_STR)
    print(f"✓ Documento compilado con éxito: {out_path}")


if __name__ == "__main__":
    out_file = OUTPUT_DIR / "PULSO_a_la_IA_Edicion_Especial.docx"
    generate_document(out_file)
